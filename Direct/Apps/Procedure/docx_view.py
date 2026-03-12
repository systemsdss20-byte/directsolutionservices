import base64
import datetime
import io

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.shortcuts import render
from django.views.generic.base import View
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docxtpl import DocxTemplate

from ... import settings
from .models import Customers, Drivers


@login_required(login_url='Procedure:login')
def generate_cover(request, idcustomer):
    if request.method == 'GET':
        try:
            path_template = settings.TEMPLATES_PDF + '/COVER.docx'
            document = Document(path_template)
            paragraphs = document.paragraphs
            obj_styles = document.styles
            obj_charstyle = obj_styles.add_style('Title1', WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(20)
            obj_font.name = 'Times New Roman'
            customer = Customers.objects.get(idcustomer=idcustomer)
            filename = '%s - %s' % (customer.idcustomer, customer.cusname)
            items = []
            if customer.dotid:
                items.append({'name': 'DOT', 'value': customer.dotid})
            if customer.boe and customer.boe != '" "':
                items.append({'name': 'UCR', 'value': customer.mcexp.strftime('%Y')})
            if customer.new_jersey and customer.new_jersey != '" "':
                items.append({'name': 'NEW JERSEY', 'value': customer.new_jersey})
            if customer.kyuid:
                items.append({'name': 'KENTUCKY', 'value': customer.kyuid})
            if customer.nyid:
                items.append({'name': 'NEW YORK', 'value': customer.nyid})
            paragraphs[-1].add_run('Customer:', style='Title1').bold = True
            paragraph = document.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format.line_spacing_rule = 2
            paragraph.add_run(filename, style='Title1')
            table = document.add_table(rows=1, cols=2)
            for item in items:
                cells = table.add_row().cells
                cells[0].add_paragraph().add_run(item['name']).bold = True
                cells[1].add_paragraph().add_run(item['value'])
            table.cell(0, 0).merge(table.cell(1, 0))
            table.cell(0, 1).merge(table.cell(1, 1))
            document.add_paragraph().add_run().add_break()
            document_data = io.BytesIO()
            document.save(document_data)
            document_data.seek(0)
            response = HttpResponse(document_data.getvalue(),
                                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = "attachment; filename=%s.docx" % filename
            response['Content-Encoding'] = "UTF-8"
            return response
        except Exception as e:
            print(e)
            return HttpResponseRedirect("/Procedure/customer_view/%s" % idcustomer)


class Certificate_random_test(View):

    def get(self, request, *args, **kwargs):
        customer = Customers.objects.only('idcustomer', 'cusname').get(pk=kwargs['customer_id'])
        return render(request, 'Procedure/Customers/Applications/DOT/certificate_random_test.html', {
            'customer': customer, 'title_certificate_random': 'Certificate Drivers Random Test',
            'urlcertificate': 'certificate_random_test'
        })

    def post(self, request, *args, **kwargs):
        customer_id = kwargs['customer_id']
        if request.method == 'POST':
            customer = Customers.objects.get(idcustomer=customer_id)
            drivers = Drivers.objects.filter(idcustomer=customer_id, random_test=True, status='Active')
            filename = '%s_RANDOM_POOL.docx' % (customer.cusname.replace(" ","_"))
            path_template = settings.TEMPLATES_PDF + '/CERTIFICADO_DRIVERS_RANDOM.docx'
            doc = DocxTemplate(path_template)
            effective_date = datetime.datetime.strptime(request.POST.get('effective_date'), '%m/%d/%Y').date()
            context = {'company_name': customer.cusname, 'dot_number': customer.dotid, 'effective_date': effective_date.strftime('%B %d, %Y'), 'drivers': drivers}
            doc.render(context)
            directory = settings.FILES_PDF + '/' + filename
            doc.save(directory)
            return JsonResponse({'filename': filename})


class Certificate_Alcohol_Drugs(View):

    def get(self, request, *args, **kwargs):
        customer = Customers.objects.only('idcustomer', 'cusname').get(pk=kwargs['customer_id'])
        return render(request, 'Procedure/Customers/Applications/DOT/certificate_alcohol_drugs.html', { 'customer': customer })

    def post(self, request, *args, **kwargs):
        customer_id = kwargs['customer_id']
        if request.method == 'POST':
            select_date = request.POST.get('select_date')
            customer = Customers.objects.get(idcustomer=customer_id)
            filename = '%s - %s ALCOHOL_DRUG.docx' % (customer.idcustomer, customer.cusname)
            path_template = settings.TEMPLATES_PDF + '/CERTIFICADO_ALCOHOL_DRUG.docx'
            doc = DocxTemplate(path_template)
            context = {'customer': customer.cusname, 'lote': customer.dotid, 'date': select_date, 'result': request.POST.get('result')}
            doc.render(context)
            directory = settings.FILES_PDF+'/'+filename
            doc.save(directory)
            return JsonResponse({'filename': filename})


class Certificate_Alcohol(View):

    def get(self, request, *args, **kwargs):
        customer = Customers.objects.only('idcustomer', 'cusname').get(pk=kwargs['customer_id'])
        drivers = Drivers.objects.filter(status='Active', idcustomer=customer, random_test=True)
        return render(request, 'Procedure/Customers/Applications/DOT/certificate_alcohol.html', { 'customer': customer, 'drivers': drivers })

    def post(self, request, *args, **kwargs):
        customer_id = kwargs['customer_id']
        if request.method == 'POST':
            select_date = datetime.datetime.strptime(request.POST.get('select_date'), '%m/%d/%Y')
            driver_id = int(request.POST.get('drivers'))
            customer = Customers.objects.get(idcustomer=customer_id)
            driver = Drivers.objects.get(iddriver=driver_id)
            filename = '%s - %s ALCOHOL.docx' % (customer.idcustomer, customer.cusname)
            path_template = settings.TEMPLATES_PDF + '/CERTIFICADO_ALCOHOL.docx'
            doc = DocxTemplate(path_template)
            context = {
                'customer': customer.cusname, 'lote': customer.dotid, 'date': select_date.strftime('%b %d, %Y'), 'result': request.POST.get('result'),
                'driver': driver.nombre
            }
            doc.render(context)
            directory = settings.FILES_PDF+'/'+filename
            doc.save(directory)
            return JsonResponse({'filename': filename})